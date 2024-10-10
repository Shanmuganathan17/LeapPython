import requests
import sys
import json
from urllib.parse import quote_plus, urlparse, urljoin
from datetime import datetime, timezone, date, timedelta
import logging as logger
import urllib3
import os
import math
import urllib
import base64
import string
import random
import re
import pytz
import string
import docx2txt  # importing doc package
import re  # importing regex package
import pandas as pd  # importing pandas package for excel
# import exceptions
# import builtins as exceptions
from docx.api import Document
import docx
from zipfile import ZipFile
import pyodbc
import warnings
import base64

warnings.filterwarnings('ignore')
# from OpenSSL import crypto
from Crypto.PublicKey import RSA
from Crypto.Signature import PKCS1_v1_5
from Crypto.Hash import SHA1
# import pyodbc
import mysql.connector
import logging as logger
# import sys
import traceback
# from leap.utils import vault
# import datetime
# import urllib3
import warnings
# import os
# from zeep import Client
import base64
import subprocess
import xmltodict

global df_new
global df_sheet_index_1
global df_sheet_index
global unitsForCaseCreation
global df_DigtalOffer_Index

urllib3.disable_warnings()
# logger.basicConfig(level=print, format='%(asctime)s INFO %(message)s', datefmt='%y/%m/%d %H:%M:%S')
arguments = sys.argv
print('Arguments')
print(arguments)

params = {}
# params = {['C:\\LEAP\\BIZOPS\\uploads\\code\\native\\APVVMSTR90963_APVendor_Master.py', 'Last_Successful_Execution_Time:Thu, 26 May 2022 13:48:01', 'Emails:{'queries':[],'id':894,'name':'APVGTMLS30347','description':'GetEmails','attributes':'{\\'bodyType\\':\\'Text\\',\\'Cacheable\\':false,\\'transformData\\':false,\\'RequestMethod\\':\\'GET\\',\\'TransformationScript\\':\\'\\',\\'Headers\\':\\'\\',\\'bodyOption\\':\\'raw\\',\\'QueryParams\\':[{\\'value\\':\\'10\\',\\'key\\':\\'NumberOfRecords\\'}],\\'Body\\':\\'\\',\\'Url\\':\\'\\/api\\/Message\\/Get\\'}','type':'r','datasource':{'id':70,'name':'APVSPNPV27517','description':'SPINE APVM','type':'REST','connectionDetails':'{\\'NoProxy\\':\\'true\\',\\'ConnectionType\\':\\'ApiRequest\\',\\'testDataset\\':{\\'name\\':\\'\\',\\'attributes\\':{\\'Endpoint\\':\\'\\',\\'RequestMethod\\':\\'GET\\',\\'Headers\\':\\'\\',\\'QueryParams\\':\\'\\',\\'Body\\':\\'\\'}},\\'AuthDetails\\':{\\'password\\':\\'\\',\\'authParams\\':{\\'grant_type\\':\\'\\',\\'client_secret\\':\\'\\',\\'client_id\\':\\'\\'},\\'authToken\\':\\'\\'},\\'AuthType\\':\\'NoAuth\\',\\'Url\\':\\'https:\\/\\/vinfngfsb-02\\',\\'fileId\\':\\'\\'}','organization':'APVendor_Master','dshashcode':'be78c8eea8592452d2cfd8b9a8953d0cd2d88114667579be00be29182718394d','activetime':'May 17, 2022, 3:32:54 PM','category':'REST','lastmodifiedby':'shanmuganathan','lastmodifieddate':'May 17, 2022, 3:32:54 PM','alias':'SPINE APVM'},'organization':'APVendor_Master','expStatus':0,'isApprovalRequired':false,'isPermissionManaged':false,'isAuditRequired':false,'isInboxRequired':false,'lastmodifiedby':'superadmin','lastmodifieddate':'May 17, 2022, 3:45:48 PM','alias':'GetEmails'}', 'Authenticate:{'queries':[],'id':895,'name':'APVATHNT97322','description':'Authenticate','attributes':'{\\'bodyType\\':\\'Text\\',\\'Cacheable\\':false,\\'transformData\\':false,\\'RequestMethod\\':\\'POST\\',\\'TransformationScript\\':\\'\\',\\'Headers\\':\\'\\',\\'bodyOption\\':\\'raw\\',\\'QueryParams\\':\\'\\',\\'Body\\':\\'eyJ1c2VybmFtZSI6ImFkbWluIiwicGFzc3dvcmQiOiJJbmZ5QDEyMyJ9\\',\\'Url\\':\\'\\/api\\/authenticate\\'}','type':'r','datasource':{'id':69,'name':'APVLPVMH97837','description':'LEAP APVM','type':'REST','connectionDetails':'{\\'NoProxy\\':\\'true\\',\\'ConnectionType\\':\\'ApiRequest\\',\\'testDataset\\':{\\'name\\':\\'\\',\\'attributes\\':{\\'Endpoint\\':\\'\\',\\'RequestMethod\\':\\'GET\\',\\'Headers\\':\\'\\',\\'QueryParams\\':\\'\\',\\'Body\\':\\'\\'}},\\'AuthDetails\\':{\\'password\\':\\'\\',\\'authParams\\':{\\'grant_type\\':\\'\\',\\'client_secret\\':\\'\\',\\'client_id\\':\\'\\'},\\'authToken\\':\\'\\'},\\'AuthType\\':\\'NoAuth\\',\\'Url\\':\\'https:\\/\\/vinfngfsb-02:8082\\',\\'fileId\\':\\'\\'}','organization':'APVendor_Master','dshashcode':'e6de14611caf2807e46de990ed990d983281eae4af17c07ce291771366a0481f','activetime':'May 17, 2022, 3:30:57 PM','category':'REST','lastmodifiedby':'shanmuganathan','lastmodifieddate':'May 17, 2022, 3:30:57 PM','alias':'LEAP APVM'},'organization':'APVendor_Master','expStatus':0,'isApprovalRequired':false,'isPermissionManaged':false,'isAuditRequired':false,'isInboxRequired':false,'lastmodifiedby':'superadmin','lastmodifieddate':'May 17, 2022, 3:51:00 PM','alias':'Authenticate'}', 'Next_Execution_Time:-1', 'projectName:apvendor_master', 'Current_Execution_Time:Thu, 26 May 2022 13:49:00', 'Last_Execution_Time:Thu, 26 May 2022 13:48:01']}
for arg in arguments:
    try:
        params[arg.split(':')[0]] = (':').join(arg.split(':')[1:])
    except:
        a = 'error'
# print(params)
proxyDict = {
    'http': '',
    'https': ''
}
pipelineStatus = 'Completed'
print(pipelineStatus)
# exit()

my_dict = {}  # to store config values

phrase = ''  # to store doc string

final_str = ''  # to store strings with checkboxes

final_dict = {}  # to store the extracted keys and values

column_list = []  # storing specific sheet column values

column_list_1 = []  # storing sheet2 specific column values

doc_Count = 0

lst_UnitsForCaseCreation = []


def getVaultConfigs(configName):
    # configParser = configparser.ConfigParser()
    # configPath = os.environ['MLSTUDIOCONFIGPATH']
    # configParser.read(configPath)
    try:
        value = os.environ[configName]
    except:
        value = ''
    return value


VAULT_URI = getVaultConfigs('VAULT_URI')
VAULT_VERSION = getVaultConfigs('VAULT_VERSION')
VAULT_VERSION = 'v1'
VAULT_APPLICATION_NAME = getVaultConfigs('VAULT_APPLICATION_NAME')
VAULT_PROFILES = getVaultConfigs('VAULT_PROFILES')
VAULT_APPROLE_ROLEID = getVaultConfigs('VAULT_APPROLE_ROLEID')
VAULT_APPROLE_SECRETID = getVaultConfigs('VAULT_APPROLE_SECRETID')

print('VAULT_URI = ', VAULT_URI)
print('VAULT_VERSION  = ', VAULT_VERSION)
print('VAULT_APPLICATION_NAME  = ', VAULT_APPLICATION_NAME)
print('VAULT_PROFILES  = ', VAULT_PROFILES)
print('VAULT_APPROLE_ROLEID  = ', VAULT_APPROLE_ROLEID)
print('VAULT_APPROLE_SECRETID  = ', VAULT_APPROLE_SECRETID)


# PPassword
def getPassword(vaultKey):
    # get Token
    isError = True
    while (isError):
        try:
            authUrl = urljoin(VAULT_URI, (VAULT_VERSION + '/auth/approle/login'))
            # authUrl='https://isvaulttst.ad.infosys.com/v1/auth/approle/login'
            authParams = {}
            authParams['role_id'] = VAULT_APPROLE_ROLEID
            authParams['secret_id'] = VAULT_APPROLE_SECRETID
            VAULT_HOST = urlparse(VAULT_URI).hostname
            PROXIES = {}
            PROXIES['http'] = ''
            PROXIES['https'] = ''
            response = requests.request(method='POST', url=authUrl, data=json.dumps(authParams), proxies=PROXIES,
                                        verify=False)
            token = ''
            if response.status_code == 200:
                responseJson = response.json()
                url = urljoin(VAULT_URI, (VAULT_VERSION + '/' + VAULT_APPLICATION_NAME + '/data/' + VAULT_PROFILES))
                # url = 'https://isvaulttst.ad.infosys.com/v1/test/data/leap'
                # print('url:',url)
                token = responseJson['auth']['client_token']

                header = {}
                header['X-Vault-Token'] = token

                response = requests.request(method='GET', url=url, headers=header, verify=False, proxies=PROXIES)
                if (response.status_code == 200):
                    try:
                        isError = False
                        resJson = response.json()
                        for key in resJson['data']['data']:
                            if key == vaultKey:
                                return resJson['data']['data'][key]

                    except:
                        logger.error('Error while retieving key from Vault')
                else:
                    logger.error(
                        'Error while retieving key from Vault. Response Code : {0}'.format(response.status_code))
            else:
                logger.error('Token Failure')
        except:
            print('some error occured in getting key value pairs')
            traceback.print_exc()


# ---------------------------------------------------------------


warnings.filterwarnings('ignore')
logger.basicConfig(level=logger.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

logger.info('SOP Classification Execution start time :' + str(datetime.now()))

os.environ['HTTP_PROXY'] = '*'
os.environ['HTTPS_PROXY'] = '*'
os.environ['NO_PROXY'] = '*'

timeStampBeforeDBConnection = datetime.now()
print('Time Stamp before DB Connection  ', str(timeStampBeforeDBConnection))

try:
    leap_username = getPassword('IS.leap.mssql.ISLEAP.userid')
    leap_password = getPassword('IS.leap.mssql.ISLEAP.password')

    leap_database = 'leapmaster_ref_data1'
    leap_server = 'ISCLSPDTDBTST\INST5'

    # leap_connection = mysql.connector.connect(user=leap_username, password=leap_password, host=leap_host,port=leap_port, database=leap_database)

    leap_connstring = 'DRIVER={/opt/microsoft/msodbcsql17/lib64/libmsodbcsql-17.10.so.1.1};Server=' + leap_server + ';DATABASE=' + leap_database + ';UID=CFG_ServiceRobo' + ';PWD=' + leap_password + ';Encrypt=yes;TrustServerCertificate=Yes;'
    leap_connection = pyodbc.connect(leap_connstring)

    # SQL query to retrieve data from a specific table
    sql_query = 'SELECT * FROM LateralOfferUpdatedConfig'

    # Using pandas to read data from SQL into a DataFrame
    df = pd.read_sql_query(sql_query, leap_connection)
    df_new = df[['name', 'value']]
    df_sheet_index_1 = df[['role', 'jobLevel', 'category', 'pUCode', 'unit']]
    df_sheet_index = df[['roleMapping', 'unitMapping', 'workLocation']]
    df_DigtalOffer_Index = df[['dtUnit', 'dtPuCode', 'dtDigitalTag', 'dtDigitalOffer']]
    my_dict = pd.Series(df_new.value.values, index=df_new.name).to_dict()
    logger.info(df_new)

    logger.info('Config sequence - completed')

    # leap_cursor = leap_connection.cursor()
    logger.info('Leap DB connection successful')

    # urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
except Exception as e:
    logger.error('Unable to connect to leapdb')
    traceback.print_exc()
    sys.exit(1)
finally:
    if leap_connection:
        leap_connection.close()
        # leap_cursor.close()
        logger.info('connection closed')

timeStampafterDBConnection = datetime.now()
print('Time Stamp after DB Connection  ', str(timeStampafterDBConnection))
print('Time taken for db connection ', str(timeStampafterDBConnection - timeStampBeforeDBConnection))


def ODS_Extraction(my_dict):
    timeStampBeforeODS_DataExtraction = datetime.now()
    print('Time Stamp before ODS Data Extraction', str(timeStampBeforeODS_DataExtraction))
    # logger.info(my_dict)
    # final_dict={}
    logger.info('Units for cse creation value fetch started')
    try:
        no_Units = my_dict['UnitsForCaseCreation']
        global lst_UnitsForCaseCreation
        lst_UnitsForCaseCreation = no_Units.split(',')
        print(lst_UnitsForCaseCreation)

    except Exception as e:

        logger.error('ERROR in reading case creation value fetch : ' + str(e))

    logger.info('ODS Data Extracting - Started')

    try:

        logger.info('Reading doc and mapping file - Started')

        # df_sheet_index_1 = pd.read_excel(my_dict['Mapping File Path'],sheet_name=my_dict['Mapping Sheet 2'])  # reading excel sheet2

        # df_sheet_index = pd.read_excel(my_dict['Mapping File Path'],sheet_name=my_dict['Mapping Sheet 1'])  # reading excel sheet1

        my_text = docx2txt.process(my_dict['Doc File Path'])  # reading ODS file as a string
        # print('docx2txt' + my_text)

        document = Document(my_dict['Doc File Path'])  # reading ODS file as a tables format
        # print('file ppath' + document)

        phrase = re.sub(r'\s+', ' ', my_text)  # converting whole multiple lines in to one single line without \n spaces

        logger.info('Reading doc and mapping file - Completed')

    except Exception as e:

        logger.error('ERROR in doc and mapping file : ' + str(e))

    my_check = True

    # Columns check in mapping file
    try:
        print('maapping ttable ccheck')
        print(my_dict['Role Mapping Column'])
        print(df_sheet_index)

        if my_dict['Role Mapping Column'] in df_sheet_index.iat[1, 1] and my_dict['Unit Mapping Column'] in \
                df_sheet_index.iat[

                    1, 2] and my_dict['Work Location Column'] in df_sheet_index.iat[1, 3] and my_dict[
            'Unnamed 1 Column'] in df_sheet_index.columns[

            1] and my_dict['Unnamed 2 Column'] in df_sheet_index.columns[2] and my_dict['Unnamed 3 Column'] in \
                df_sheet_index.columns[3]:
            logger.info('Columns check passed in mapping file sheet1 : ' + my_dict['Mapping Sheet 1'])

        if my_dict['Role Column'] in df_sheet_index_1.columns[0] and my_dict['Job Level Column'] in \
                df_sheet_index_1.columns[1] and my_dict['Category Column'] in df_sheet_index_1.columns[2]:

            logger.info('Columns check passed in mapping file sheet2 : ' + my_dict['Mapping Sheet 2'])

        else:

            my_check = False

            logger.error('Columns check failed in mapping file')

    except Exception as e:

        logger.error('ERROR in Columns check in mapping file ' + str(e))

    # String checks in doc string

    try:
        print(my_dict)
        print(phrase)

        if my_dict['Candidate ID ODS Field'] in phrase and my_dict['Indent number ODS'] in phrase and my_dict[
            'NAME IN FULL ODS'] in phrase and my_dict['iRACE Role Designation ODS'] in phrase and my_dict[
            'JL JSL ODS'] in phrase and my_dict['Allocated to Unit ODS'] in phrase and my_dict[
            'Location of Posting ODS'] in phrase and my_dict['Relocation Allowance ODS'] in phrase and my_dict[
            'SEZ ODS'] in phrase and my_dict['STP ODS'] in phrase and my_dict['Skill Group ODS'] in phrase and my_dict[
            'NSO ODS'] in phrase and my_dict['Work From Home ODS'] in phrase and my_dict[
            'Remote Work Location ODS'] in phrase and my_dict['End Date ODS'] in phrase and my_dict[
            'Technology Expertise ODS'] in phrase and my_dict['Highest Qualification ODS'] in phrase and my_dict[
            'Experience ODS'] in phrase and my_dict['PG ODS'] in phrase and my_dict['Engineering ODS'] in phrase and \
                my_dict['Non Engineering ODS'] in phrase and my_dict['DOJ ODS'] in phrase and my_dict[
            'Salary Proposed ODS'] in phrase and my_dict['Role Max ODS'] in phrase and my_dict[
            'Any Exceptions ODS'] in phrase and my_dict['Joining bonus ODS'] in phrase and my_dict[
            'T50 ODS'] in phrase and my_dict['AE ODS'] in phrase and my_dict['Stretch ODS'] in phrase and my_dict[
            'CC Offer Softcopy to ODS'] in phrase and my_dict['Panelist ODS'] in phrase and my_dict[
            'Recruiter Sign Off ODS'] in phrase and my_dict['Recruitment Manager ODS'] in phrase and my_dict[
            'Candidate Name ODS'] in phrase and my_dict['Account Name ODS'] in phrase and my_dict[
            'Source ODS'] in phrase and my_dict['Source Emp Id for ConnectInfy ODS'] in phrase and my_dict[
            'Restart With Infosys ODS'] in phrase and my_dict['Subcon ID for subcon conversion cases ODS'] in phrase and \
                my_dict['Panelist Details ODS'] in phrase:

            logger.info('String checks passed in ODS Document ')

        else:

            my_check = False

            logger.error('String checks failed in ODS Document ')

    except Exception as e:

        logger.error('ERROR in String checks in ODS Document ' + str(e))

    def find_between_r(s, first, last):

        try:

            start = s.index(first) + len(first)

            end = s.index(last, start)

            return s[start:end]

        except ValueError:

            return ''

    # finding the uncheckboxes and returning thr selected checkbox string

    def check_boxes(between_str):

        try:

            list_str = str(between_str).split(' ')

            output_str = ''

            for index in range(len(list_str) - 1):

                if list_str[index] != '\u2610' or not (str('\u2610') in list_str[index]):

                    if list_str[index + 1].isalpha():
                        output_str = list_str[index + 1]

            if len(output_str) <= 0:
                output_str = list_str[0]

            if not output_str == '\u2610' or (output_str == str('\u2610')):
                return (output_str)

        except Exception as e:

            logger.error('ERROR in check_boxes function : ' + str(e))

    # returning specific column and sheet values as a list

    def ColumnsData(df_sheet_index, sheetname):

        try:

            df_sheet_index.dropna(subset=[sheetname], inplace=True)

            return (df_sheet_index[sheetname].tolist())

        except Exception as e:

            logger.error('ERROR in ColumnsData function : ' + str(e))

    if my_check:
        # Candidate Name criteria check

        try:

            if str(find_between_r(phrase, my_dict['Candidate Name ODS'],
                                  my_dict['iRACE Role Designation ODS'])).strip() is not None:

                final_dict.update(

                    {'Candidate Name': str(
                        find_between_r(phrase, my_dict['Candidate Name ODS'],
                                       my_dict['iRACE Role Designation ODS'])).strip()})

            else:

                logger.error('Candidate Name is not equal to criteria of alphabet')

                logger.error('Candidate Name - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Candidate Name ODS'],
                                   my_dict['iRACE Role Designation ODS'])).strip())

                final_dict.update({'Candidate Name': None})

        except Exception as e:

            logger.error('ERROR in extracting Candidate Name : ' + str(e))

            final_dict.update({'Candidate Name': None})
        # Account Name criteria check whether it is a alpha or not

        try:

            if str(find_between_r(phrase, my_dict['Account Name ODS'],
                                  my_dict['Location of Posting ODS'])).strip() is not None:

                final_dict.update(

                    {'Account Name': str(
                        find_between_r(phrase, my_dict['Account Name ODS'],
                                       my_dict['Location of Posting ODS'])).strip()})

            else:

                logger.error('Account Name is not equal to criteria of alphabet')

                logger.error('Account Name - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Account Name ODS'], my_dict['Location of Posting ODS'])).strip())

                final_dict.update({'Account Name': None})

        except Exception as e:

            logger.error('ERROR in extracting Account Name : ' + str(e))

            final_dict.update({'Account Name': None})

        # Source criteria check whether it is a alpha or not

        try:

            if str(find_between_r(phrase, my_dict['Source ODS'],
                                  my_dict['Source Emp Id for ConnectInfy ODS'])).strip() is not None:

                final_dict.update(

                    {'Source': str(
                        find_between_r(phrase, my_dict['Source ODS'],
                                       my_dict['Source Emp Id for ConnectInfy ODS'])).strip()})

            else:

                logger.error('Source is not equal to criteria of alphabet')

                logger.error('Source - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Source ODS'],
                                   my_dict['Source Emp Id for ConnectInfy ODS'])).strip())

                final_dict.update({'Source': None})

        except Exception as e:

            logger.error('ERROR in extracting Source : ' + str(e))

            final_dict.update({'Source': None})

        # Technology criteria check

        try:

            if str(find_between_r(phrase, my_dict['Technology Expertise ODS'],
                                  my_dict['Skill Group ODS'])).strip() is not None:

                final_dict.update(

                    {'Technology': str(
                        find_between_r(phrase, my_dict['Technology Expertise ODS'],
                                       my_dict['Skill Group ODS'])).strip()})

            else:

                logger.error('Technology is not equal to criteria of alphabet')

                logger.error('Technology - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Technology Expertise ODS'],
                                   my_dict['Skill Group ODS'])).strip())

                final_dict.update({'Technology': None})

        except Exception as e:

            logger.error('ERROR in extracting Source : ' + str(e))

            final_dict.update({'Technology': None})
        # Restart With Infosys criteria check
        try:

            final_str = str(find_between_r(phrase, my_dict['Restart With Infosys ODS'],
                                           my_dict['Subcon ID for subcon conversion cases ODS'])).strip()

            if final_str.count('\u2610') == 1:

                final_dict.update({'Restart With Infosys': check_boxes(final_str)})

            else:

                logger.error('Restart With Infosys - criteria check: Failed ' + final_str)

                final_dict.update({'Restart With Infosys': None})

        except Exception as e:

            logger.error('ERROR in extracting Restart With Infosys : ' + str(e))

            final_dict.update({'Restart With Infosys': None})
        # candidate id criteria check up to 10 numeric digits

        try:

            if len(str(find_between_r(phrase, my_dict['Candidate ID ODS Field'],
                                      my_dict['Indent number ODS'])).strip()) == 10 and (

                    str(find_between_r(phrase, my_dict['Candidate ID ODS Field'],
                                       my_dict['Indent number ODS'])).strip().isnumeric()):

                final_dict.update(

                    {'Candidate ID': str(find_between_r(phrase, my_dict['Candidate ID ODS Field'],
                                                        my_dict['Indent number ODS'])).strip()})

            else:

                logger.error('Candidate ID is not 10-digit number or alphanumeric value or special character')

                logger.error('Candidate ID - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Candidate ID ODS Field'], my_dict['Indent number ODS'])).strip())

                final_dict.update({'Candidate ID': None})

        except Exception as e:

            logger.error('ERROR in extracting Candidate ID : ' + str(e))

            final_dict.update({'Candidate ID': None})

        # indent number criteria check whether it is a alphanumeric or not

        try:

            if str(find_between_r(phrase, my_dict['Indent number ODS'], my_dict['NAME IN FULL ODS'])).strip().isalnum():

                final_dict.update(

                    {'Indent': str(
                        find_between_r(phrase, my_dict['Indent number ODS'], my_dict['NAME IN FULL ODS'])).strip()})

            else:

                logger.error('Indent number (if applicable) is not equal to criteria of alphanumeric')

                logger.error('Indent number - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Indent number ODS'], my_dict['NAME IN FULL ODS'])).strip())

                final_dict.update({'Indent': None})

        except Exception as e:

            logger.error('ERROR in extracting Indent : ' + str(e))

            final_dict.update({'Indent': None})
        # NSO logic check
        try:

            if str(find_between_r(phrase, my_dict['NSO ODS'], my_dict['Source ODS'])).strip() is not None:

                final_dict.update(

                    {'NSO': str(
                        find_between_r(phrase, my_dict['NSO ODS'], my_dict['Source ODS'])).strip()})

            else:

                logger.error('NSO (For Digital / Adjacent Skill Group) is not equal to criteria of alpha')

                logger.error('NSO (For Digital / Adjacent Skill Group) - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['NSO ODS'], my_dict['Source ODS'])).strip())

                final_dict.update({'NSO': None})

        except Exception as e:

            logger.error('ERROR in extracting NSO : ' + str(e))

            final_dict.update({'NSO': None})

        # iRACE role designation criteria check if that string is mapped into the Role Mapping column or not

        try:

            column_list = ColumnsData(df_sheet_index, my_dict['Unnamed 1 Column'])

            for i in range(len(column_list)):

                if str(find_between_r(phrase, my_dict['iRACE Role Designation ODS'],

                                      my_dict['JL JSL ODS'])).strip().lower() == str(column_list[i]).lower():

                    final_dict.update({'Role Designation': str(

                        find_between_r(phrase, my_dict['iRACE Role Designation ODS'], my_dict['JL JSL ODS'])).strip()})

                    break

                else:

                    final_dict.update({'Role Designation': None})

            if final_dict['Role Designation'] != None:

                # updating Job level and Category based on the mapping Excel sheet (roles in role mapping)

                column_list_1 = ColumnsData(df_sheet_index_1, my_dict['Role Column'])

                for j in range(len(column_list_1)):

                    if str(find_between_r(phrase, my_dict['iRACE Role Designation ODS'],

                                          my_dict['JL JSL ODS'])).strip().lower() == str(column_list_1[j]).lower():

                        al = ''

                        jl = ''

                        num = [int(x) for x in df_sheet_index_1.iloc[j][my_dict['Job Level Column']] if x.isdigit()]

                        for l in num:
                            jl = jl + str(l);

                        alpha = [(x) for x in df_sheet_index_1.iloc[j][my_dict['Job Level Column']] if x.isalpha()]

                        for a in alpha:
                            al = al + str(a);

                        final_dict.update({'Job Level': jl})

                        final_dict.update({'personal level': jl})

                        final_dict.update({'Job sub-Level': al})

                        final_dict.update({'personal sub-Level': al})

                        final_dict.update({'Category': df_sheet_index_1.iloc[j][my_dict['Category Column']]})

                        break

            else:

                logger.info('Role Destination is not correctly updated according to criteria check - ' + str(

                    find_between_r(phrase, my_dict['iRACE Role Designation ODS'], my_dict['JL JSL ODS'])).strip())

                final_dict.update({'Job Level': None})

                final_dict.update({'personal level': None})

                final_dict.update({'Job sub-Level': None})

                final_dict.update({'personal sub-Level': None})

                final_dict.update({'Category': None})

        except Exception as e:

            logger.info('ERROR in extracting iRACE role designation : ' + str(e))

            final_dict.update({'Role Designation': None})

            final_dict.update({'Job Level': None})

            final_dict.update({'personal level': None})

            final_dict.update({'Job sub-Level': None})

            final_dict.update({'personal sub-Level': None})

            final_dict.update({'Category': None})

        # Allocate to (Unit) criteria check if that string is mapped into the Unit Mapping column or not

        try:

            column_list = ColumnsData(df_sheet_index_1, my_dict['Unnamed 2 Column'])

            for i in range(len(column_list)):

                if str(find_between_r(phrase, my_dict['Allocated to Unit ODS'],
                                      my_dict['Account Name ODS'])).strip().lower() == str(

                    column_list[i]).lower():

                    final_dict.update({'Practice Unit': str(

                        find_between_r(phrase, my_dict['Allocated to Unit ODS'], my_dict['Account Name ODS'])).strip()})
                    logger.info('Unit : ' + str(df_sheet_index_1.iloc[i][my_dict['pUCode Column']]))

                    final_dict.update({'Unit': df_sheet_index_1.iloc[i][my_dict['pUCode Column']]})
                    # final_dict.update({'Unit':'BEF'})

                    break

                else:

                    final_dict.update({'Practice Unit': None})
                    final_dict.update({'Unit': None})

            if final_dict['Practice Unit'] == None:
                logger.error('Allocate to (Unit) criteria check : Failed ' + str(

                    find_between_r(phrase, my_dict['Allocated to Unit ODS'],
                                   my_dict['Location of Posting ODS'])).strip())

        except Exception as e:

            logger.error('ERROR in extracting Allocate to (Unit) : ' + str(e))

            final_dict.update({'Practice Unit': None})
            final_dict.update({'Unit': None})

        # Location of Posting  criteria check if that string is mapped into the Work Location column or not

        try:

            column_list = ColumnsData(df_sheet_index, my_dict['Unnamed 3 Column'])

            for i in range(len(column_list)):

                if str(find_between_r(phrase, my_dict['Location of Posting ODS'], '\u2610')).strip().split(' ')[
                    0].lower() == str(

                    column_list[i]).lower():

                    final_dict.update({'Location of Posting':

                                           str(find_between_r(phrase, my_dict['Location of Posting ODS'],
                                                              '\u2610')).strip().split(

                                               ' ')[

                                               0]})

                    break

                else:

                    final_dict.update({'Location of Posting': None})

            if final_dict['Location of Posting'] == None:
                logger.error('Location of Posting criteria check : Failed ' +

                             str(find_between_r(phrase, my_dict['Location of Posting ODS'], '\u2610')).strip().split(
                                 ' ')[0])

        except Exception as e:

            logger.error('ERROR in extracting Location of Posting : ' + str(e))

            final_dict.update({'Location of Posting': None})

        # Establishment type criteria check

        try:

            final_str = (str(find_between_r(phrase, my_dict['Location of Posting ODS'],
                                            my_dict['Relocation Allowance ODS'])).strip())

            if any(chr.isdigit() for chr in final_str) and final_str.count('\u2610') == 1 and int(

                    str(find_between_r(phrase, my_dict['SEZ ODS'], my_dict['Relocation Allowance ODS'])).replace('_',
                                                                                                                 '').strip()) >= 1 and int(

                str(find_between_r(phrase, my_dict['SEZ ODS'], my_dict['Relocation Allowance ODS'])).replace('_',
                                                                                                             '').strip()) <= 10:

                final_dict.update({'Establishment': my_dict['SEZ ODS'] + str(

                    find_between_r(phrase, my_dict['SEZ ODS'], my_dict['Relocation Allowance ODS'])).replace('_',
                                                                                                             '').strip()})

            # elif final_str.count('?') == 1 and check_boxes(final_str).upper() == 'STP':

            elif not (any(chr.isdigit() for chr in final_str)) and my_dict['STP ODS'] in (
                    final_str).upper() and final_str.count(

                '\u2610') == 0:

                final_dict.update({'Establishment': my_dict['STP ODS']})

            else:

                logger.error('Establishment (Location of Posting) - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Location of Posting ODS'],
                                   my_dict['Relocation Allowance ODS'])).strip())

                final_dict.update({'Establishment': None})

        except Exception as e:

            logger.error('ERROR in extracting Location of Posting(Establishment) : ' + str(e))

            final_dict.update({'Establishment': None})

        # Skill group criteria check

        try:

            final_str = str(find_between_r(phrase, my_dict['Skill Group ODS'], my_dict['NSO ODS'])).strip()

            if final_str.count('\u2610') == 2:

                final_dict.update({'Hiring Type': check_boxes(final_str)})

            else:

                logger.error('Skill Group - criteria check: Failed ' + final_str)

                final_dict.update({'Hiring Type': None})

        except Exception as e:

            logger.error('ERROR in extracting Skill Group : ' + str(e))

            final_dict.update({'Hiring Type': None})

        # Work from home criteria check

        try:

            final_str = str(
                find_between_r(phrase, my_dict['Work From Home ODS'], my_dict['Remote Work Location ODS'])).strip()

            if str(check_boxes(final_str)).upper() == 'NA' and final_str.count('\u2610') == 2:

                final_dict.update({'Work From Home': check_boxes(final_str)})

                final_dict.update({'Remote work from home': 'No'})

                final_dict.update({'Permanent work from home': 'No'})

                final_dict.update({'Remote work Location': None})
                final_dict.update({'partTimeEmployee': None})

                final_dict.update({'End Date': None})

            elif str(check_boxes(final_str)).upper() == 'PERMANENT' and final_str.count('\u2610') == 2:

                final_dict.update({'Work From Home': check_boxes(final_str)})

                final_dict.update({'Remote work from home': 'No'})

                final_dict.update({'Permanent work from home': 'Yes'})
                final_dict.update({'partTimeEmployee': None})

                # final_dict.update({'Remote work Location': None})

                final_dict.update({'End Date': None})
                column_list = []

                column_list = ColumnsData(df_sheet_index, my_dict['Unnamed 3 Column'])

                for i in range(len(column_list)):

                    if (

                            str(find_between_r(phrase, my_dict['Remote Work Location ODS'],
                                               my_dict['End Date ODS'])).strip()) == column_list[i]:
                        final_dict.update({'Remote work Location': str(

                            find_between_r(phrase, my_dict['Remote Work Location ODS'],
                                           my_dict['End Date ODS'])).strip()})

                        break

            elif str(check_boxes(final_str)).upper() == 'REMOTE' and final_str.count('\u2610') == 2:

                final_dict.update({'Work From Home': check_boxes(final_str)})

                final_dict.update({'Remote work from home': 'Yes'})

                final_dict.update({'Permanent work from home': 'No'})
                final_dict.update({'partTimeEmployee': 'No'})

                column_list = []

                column_list = ColumnsData(df_sheet_index, my_dict['Unnamed 3 Column'])

                for i in range(len(column_list)):

                    if (

                            str(find_between_r(phrase, my_dict['Remote Work Location ODS'],
                                               my_dict['End Date ODS'])).strip()) == column_list[i]:

                        final_dict.update({'Remote work Location': str(

                            find_between_r(phrase, my_dict['Remote Work Location ODS'],
                                           my_dict['End Date ODS'])).strip()})

                        break

                    else:

                        final_dict.update({'Remote work Location': None})

                if final_dict['Remote work Location'] == None:
                    logger.error('Work from home - remote work location criteria check - failed ' + str(

                        find_between_r(phrase, my_dict['Remote Work Location ODS'], my_dict['End Date ODS'])).strip())

                try:

                    date_object = datetime.strptime(

                        str(find_between_r(phrase, my_dict['End Date ODS'],
                                           my_dict['Technology Expertise ODS'])).strip(),

                        '%m-%d-%Y').date()

                    # print(type(date_object))

                    # print(date_object)

                    final_dict.update(

                        {'End Date': str(

                            find_between_r(phrase, my_dict['End Date ODS'],
                                           my_dict['Technology Expertise ODS'])).strip()})

                except Exception as e:

                    logger.error('Error in fetching end date format : ' + str(e))

                    final_dict.update({'End Date': None})

            else:

                logger.error('Work From Home - criteria check: Failed ' + final_str)

                final_dict.update({'Work From Home': None})

                final_dict.update({'Remote work from home': None})

                final_dict.update({'Permanent work from home': None})

                final_dict.update({'Remote work Location': None})

                final_dict.update({'End Date': None})
                final_dict.update({'partTimeEmployee': None})

        except Exception as e:

            logger.error('ERROR in extracting Work From Home : ' + str(e))

            final_dict.update({'Work From Home': None})

            final_dict.update({'Remote work from home': None})

            final_dict.update({'Permanent work from home': None})

            final_dict.update({'Remote work Location': None})

            final_dict.update({'End Date': None})
            final_dict.update({'partTimeEmployee': None})

        # Highest qualification criteria check

        try:

            final_str = str(
                find_between_r(phrase, my_dict['Highest Qualification ODS'], my_dict['Experience ODS'])).strip()

            if '(' in final_str or ')' in final_str or '/' in final_str:
                final_str = final_str.replace(' (', '').replace('2)', '').replace('/', '').replace('3)', '').replace(

                    ' 1/',

                    '').replace(

                    '-', '').replace(' 1', '')  # replacing special chars, numbers

            if final_str.count('\u2610') == 3:

                if check_boxes(final_str) in 'PGTierTier':

                    final_dict.update({'Education Category': 'PG(Tier 1/ Tier 2)'})

                elif check_boxes(final_str) in 'EngineeringPGTier':

                    final_dict.update({'Education Category': 'Engineering/ PG(Tier 3)'})

                elif check_boxes(final_str) in 'NonEngineering':

                    final_dict.update({'Education Category': 'Non Engineering'})

                else:

                    final_dict.update({'Education Category': check_boxes(final_str)})

            else:

                logger.error('Highest qualification - criteria check: Failed ' + final_str)

                final_dict.update({'Education Category': None})

        except Exception as e:

            logger.error('ERROR in extracting Highest qualification : ' + str(e))

            final_dict.update({'Education Category': None})

        # Experience criteria check

        try:

            if str(find_between_r(phrase, my_dict['Experience ODS'], my_dict['DOJ ODS'])).strip().isnumeric():

                final_dict.update({'Experience (Month)': str(
                    find_between_r(phrase, my_dict['Experience ODS'], my_dict['DOJ ODS'])).strip()})

            else:

                logger.error(

                    'Experience - criteria check: Failed ' + str(
                        find_between_r(phrase, my_dict['Experience ODS'], my_dict['DOJ ODS'])).strip())

                final_dict.update({'Experience (Month)': None})

        except Exception as e:

            logger.error('ERROR in extracting Experience : ' + str(e))

            final_dict.update({'Experience (Month)': None})

        # DOJ criteria check

        try:

            date_object = datetime.strptime(
                str(find_between_r(phrase, my_dict['DOJ ODS'], my_dict['Salary Proposed ODS'])).strip(),

                '%m-%d-%Y').date()

            final_dict.update({'Date of joining': str(
                find_between_r(phrase, my_dict['DOJ ODS'], my_dict['Salary Proposed ODS'])).strip()})

        except Exception as e:

            logger.error('DOJ - criteria check: Failed - wrong date format ' + str(e))

            final_dict.update({'Date of joining': None})

        # Salary Proposed criteria check

        try:

            if str(find_between_r(phrase, my_dict['Salary Proposed ODS'],
                                  my_dict['Role Max ODS'])).strip().upper().count('LPA') == 1:

                final_dict.update({'Salary Proposed (in Lacs)': round(

                    float((str(find_between_r(phrase, my_dict['Salary Proposed ODS'],
                                              my_dict['Role Max ODS'])).strip().split())[0]) * 100000,

                    0)})

            else:

                logger.error('Salary Proposed - criteria check: Failed ' + str(

                    find_between_r(phrase, my_dict['Salary Proposed ODS'], my_dict['Role Max ODS'])).strip())

                final_dict.update({'Salary Proposed (in Lacs)': None})

        except Exception as e:

            logger.error('ERROR in extracting Salary Proposed : ' + str(e))

            final_dict.update({'Salary Proposed (in Lacs)': None})

        # Any Exceptions criteria check

        try:

            if str(find_between_r(phrase, my_dict['Any Exceptions ODS'],
                                  my_dict['Joining bonus ODS'])).strip() != '\u2610':

                if (str(find_between_r(phrase, my_dict['Joining bonus ODS'], 'Lac')).strip()):

                    final_dict.update(

                        {'Joining Bonus': float(
                            str(find_between_r(phrase, my_dict['Joining bonus ODS'], 'Lac')).strip()) * 100000})

                elif (str(find_between_r(phrase, my_dict['Joining bonus ODS'], 'lac')).strip()):

                    final_dict.update(

                        {'Joining Bonus': float(
                            str(find_between_r(phrase, my_dict['Joining bonus ODS'], 'lac')).strip()) * 100000})

                elif (str(find_between_r(phrase, my_dict['Joining bonus ODS'], 'K')).strip()):

                    final_dict.update(

                        {'Joining Bonus': float(
                            str(find_between_r(phrase, my_dict['Joining bonus ODS'], 'K')).strip()) * 1000})

                else:

                    logger.error('Any exceptions (Joining Bonus) - criteria check: failed ' + str(

                        find_between_r(phrase, my_dict['Any Exceptions ODS'], my_dict['Joining bonus ODS'])).strip())

                    final_dict.update({'Joining Bonus': None})

            else:

                logger.error('Any exceptions (Joining Bonus) - criteria check: not selected ' + str(

                    find_between_r(phrase, my_dict['Any Exceptions ODS'], my_dict['Joining bonus ODS'])).strip())

                final_dict.update({'Joining Bonus': None})

        except Exception as e:

            logger.error('ERROR in extracting Any exceptions (Joining Bonus) : ' + str(e))

            final_dict.update({'Joining Bonus': None})

        # T50 criteria check

        try:

            if not '\u2610' in str(find_between_r(phrase, my_dict['Joining bonus ODS'], my_dict['T50 ODS'])).strip():

                final_dict.update({'T50': 'Yes'})

            else:

                logger.error('T50 - criteria check: not selected')

                final_dict.update({'T50': 'No'})

        except Exception as e:

            logger.error('ERROR in extracting T50 : ' + str(e))

            final_dict.update({'T50': None})

        # AE criteria check

        try:

            if not '\u2610' in str(find_between_r(phrase, my_dict['T50 ODS'], my_dict['AE ODS'])).strip() and not str(

                    find_between_r(phrase, my_dict['AE ODS'], '%').replace('-', '').strip()) == '':

                final_dict.update(
                    {'AE': float(str(find_between_r(phrase, my_dict['AE ODS'], '%').replace('-', '').strip()))})

            else:

                logger.error('AE - criteria check: failed ' + str(
                    find_between_r(phrase, my_dict['T50 ODS'], my_dict['AE ODS'])).strip())

                final_dict.update({'AE': None})

        except Exception as e:

            logger.error('ERROR in extracting AE : ' + str(e))

            final_dict.update({'AE': None})

        # Stretch % criteria check

        try:

            if str(find_between_r(phrase, my_dict['Stretch ODS'], my_dict['Any Exceptions ODS'])).replace('%',
                                                                                                          '').strip():

                final_dict.update(

                    {'Stretch %': round(

                        float(
                            str(find_between_r(phrase, my_dict['Stretch ODS'], my_dict['Any Exceptions ODS'])).replace(
                                '%', '').strip()), 2)})

            else:

                logger.error('Stretch % - criteria check: Failed - blank or not an integer value')

                final_dict.update({'Stretch %': None})

        except Exception as e:

            logger.error('ERROR in extracting Stretch % : ' + str(e))

            final_dict.update({'Stretch %': None})

        # CC Offer Softcopy to criteria check

        try:

            if str(find_between_r(phrase, my_dict['CC Offer Softcopy to ODS'],
                                  my_dict['Panelist ODS'])).strip().isascii():

                final_dict.update(

                    {'CC Employee Name': str(

                        find_between_r(phrase, my_dict['CC Offer Softcopy to ODS'],
                                       my_dict['Panelist ODS'])).strip()})

                # logger.info(my_dict['CC Offer Softcopy to ODS']+' : '+my_dict['Panelist Details ODS']+' : '+str(find_between_r(phrase, my_dict['CC Offer Softcopy to ODS'],my_dict['Panelist Details ODS'])).strip())

            else:

                logger.error('CC Offer Softcopy - criteria check: Failed - not a alpha string ' + str(

                    find_between_r(phrase, my_dict['CC Offer Softcopy to ODS'],
                                   my_dict['Panelist ODS'])).strip())

                final_dict.update({'CC Employee Name': None})

        except Exception as e:

            logger.error('ERROR in extracting CC Offer Softcopy : ' + str(e))

            final_dict.update({'CC Employee Name': None})

        # Panelist Details criteria check

        try:

            # table = document.tables[len(document.tables) - 1]

            table = document.tables[4]

            data = []
            keys = None

            for i, row in enumerate(table.rows):

                text = (cell.text for cell in row.cells)

                if i == 0:
                    keys = tuple(text)

                    continue

                row_data = dict(zip(keys, text))

                data.append(row_data)

            try:
                if str(data[0].get(my_dict['Panelist Details ODS'])).strip() is not None:
                    a = str(data[0].get(my_dict['Panelist Details ODS'])).strip()
                    Panelist_Details = []
                    Panelist_Details = re.split('/', a)
                    final_dict.update({'Employee Name1': Panelist_Details[0].strip()})
                    final_dict.update({'Employee ID1': Panelist_Details[1].strip()})
                    final_dict.update({'Employee Role1': Panelist_Details[2].strip()})
                else:
                    logger.error('Panelist Details1 - criteria check: Failed ' + str(
                        data[0].get(my_dict['Panelist Details ODS'])).strip())

                    final_dict.update({'Employee Name1': None})
                    final_dict.update({'Employee ID1': None})
                    final_dict.update({'Employee Role1': None})
            except Exception as e:
                logger.error('ERROR in extracting Panelist Details 1 : ' + str(e))
                final_dict.update({'Employee Name1': None})
                final_dict.update({'Employee ID1': None})
                final_dict.update({'Employee Role1': None})
            try:
                if str(data[1].get(my_dict['Panelist Details ODS'])).strip() is not None:
                    a = str(data[1].get(
                        my_dict['Panelist Details ODS'])).strip()
                    Panelist_Details = []
                    Panelist_Details = re.split('/', a)
                    final_dict.update({'Employee Name2': Panelist_Details[0].strip()})
                    final_dict.update({'Employee ID2': Panelist_Details[1].strip()})
                    final_dict.update({'Employee Role2': Panelist_Details[2].strip()})
                else:
                    logger.error('Panelist Details1 - criteria check: Failed ' + str(
                        data[1].get(my_dict['Panelist Details ODS'])).strip())

                    final_dict.update({'Employee Name2': None})
                    final_dict.update({'Employee ID2': None})
                    final_dict.update({'Employee Role2': None})
            except Exception as e:
                logger.error('ERROR in extracting Panelist Details 2 : ' + str(e))
                final_dict.update({'Employee Name2': None})
                final_dict.update({'Employee ID2': None})
                final_dict.update({'Employee Role2': None})
            try:
                if str(data[2].get(my_dict['Panelist Details ODS'])).strip() is not None:
                    a = str(data[2].get(
                        my_dict['Panelist Details ODS'])).strip()
                    Panelist_Details = []
                    Panelist_Details = re.split('/', a)
                    final_dict.update({'Employee Name3': Panelist_Details[0].strip()})
                    final_dict.update({'Employee ID3': Panelist_Details[1].strip()})
                    final_dict.update({'Employee Role3': Panelist_Details[2].strip()})
                else:
                    logger.error('Panelist Details1 - criteria check: Failed ' + str(
                        data[2].get(my_dict['Panelist Details ODS'])).strip())

                    final_dict.update({'Employee Name3': None})
                    final_dict.update({'Employee ID3': None})
                    final_dict.update({'Employee Role3': None})
            except Exception as e:
                logger.error('ERROR in extracting Panelist Details 3 : ' + str(e))
                final_dict.update({'Employee Name3': None})
                final_dict.update({'Employee ID3': None})
                final_dict.update({'Employee Role3': None})
        except Exception as e:
            logger.error('ERROR in extracting Panelist Details : ' + str(e))

        # Recruiter Sign Off and Recruitment Manager criteria check

        try:

            # table = document.tables[len(document.tables) - 1]

            table = document.tables[6]

            data = []

            keys = None

            for i, row in enumerate(table.rows):

                text = (cell.text for cell in row.cells)

                if i == 0:
                    keys = tuple(text)

                    continue

                row_data = dict(zip(keys, text))

                data.append(row_data)

            if str(data[0].get(my_dict['Recruiter Sign Off ODS'])).strip().isnumeric():

                final_dict.update(
                    {'Recruiter ID - Contact Details': str(data[0].get(my_dict['Recruiter Sign Off ODS'])).strip()})

            else:

                logger.error(

                    'Recruiter Sign Off - criteria check: Failed ' + str(
                        data[0].get(my_dict['Recruiter Sign Off ODS'])).strip())

                final_dict.update({'Recruiter ID - Contact Details': None})

            if str(data[0].get(my_dict['Recruitment Manager ODS'])).strip().isnumeric():

                final_dict.update(
                    {'Recruiter Manager ID': str(data[0].get(my_dict['Recruitment Manager ODS'])).strip()})

            else:

                logger.error(

                    'Recruitment Manager - criteria check: Failed ' + str(
                        data[0].get(my_dict['Recruitment Manager ODS'])).strip())

                final_dict.update({'Recruiter Manager ID': None})

            # Names of manager and recuriters

            if str(data[1].get(my_dict['Recruiter Sign Off ODS'])).strip().isascii():

                final_dict.update(
                    {'Recuriter Name - Contact Details': str(data[1].get(my_dict['Recruiter Sign Off ODS'])).strip()})

            else:

                logger.error(

                    'Recruiter Sign Off - criteria check: Failed ' + str(
                        data[1].get(my_dict['Recruiter Sign Off ODS'])).strip())

                final_dict.update({'Recuriter Name - Contact Details': None})

            if str(data[1].get(my_dict['Recruitment Manager ODS'])).strip().isascii():

                final_dict.update(
                    {'Recuriter Manager Name': str(data[1].get(my_dict['Recruitment Manager ODS'])).strip()})

            else:

                logger.error('Recruitment Manager - criteria check: Failed ' + str(
                    data[1].get(my_dict['Recruitment Manager ODS'])))

                final_dict.update({'Recuriter Manager Name': None})

        except Exception as e:

            logger.error('ERROR in extracting Recruiter Sign Off and Recruitment Manager : ' + str(e))

            final_dict.update({'Recruiter Manager ID': None})

            final_dict.update({'Recruiter ID - Contact Details': None})

            final_dict.update({'Recuriter Manager Name': None})

            final_dict.update({'Recuriter Name - Contact Details': None})



    else:

        logger.error('Column check or string checks - Failed ')

    logger.info('Final Dictionary - ' + str(final_dict))

    logger.info('ODS Document Extracting - Completed')

    timeStampafterODS_Extraction = datetime.now()
    print('Time Stamp after ODS Data Extraction ', str(timeStampafterODS_Extraction))
    print('Time taken for ODS Data Extraction ', str(timeStampafterODS_Extraction - timeStampBeforeODS_DataExtraction))

    return (final_dict)


def parse_attachments(attachments, mailId):
    timeStampBefore_ParseAttachment = datetime.now()
    print('Time Stamp before Parse Attachment Method  ', str(timeStampBefore_ParseAttachment))

    logger.info('pparse aattachments sstarted')
    attachmentUrl = spineUrl + '/api/Message/GetAttachment?attachmentIDs={0}'
    emlUrl = spineUrl + '/api/Message/GetEmlFiles?emailIds={0}'
    attachmentLinks = []
    global doc_Count
    doc_Count = 0
    global final_dict
    final_dict = {}
    logger.info('Attachments found: ' + str(len(attachments)))
    attachments.append({'attachmentID': 0})
    print('Attachments found: ' + str(len(attachments)))
    for attachment in attachments:
        attachmentLink = {}
        id = attachment['attachmentID']
        if (id == 0):
            emlUrl = emlUrl.format(mailId)
            response = requests.get(url=emlUrl, verify=False, proxies=proxyDict)
            if (response.status_code != 200):
                logger.error('Error fetching attachment')
                continue
        else:
            getAttachmentUrl = attachmentUrl.format(id)
            print(getAttachmentUrl)
            response = requests.get(url=getAttachmentUrl, verify=False, proxies=proxyDict)
            if (response.status_code != 200):
                logger.error('Error fetching attachment')
                continue
        print(response.json())
        attach = response.json()[0]
        filename = attach['fileName'] + attach['extension']
        file = open(filename, 'wb')
        file.write(base64.b64decode(attach['content']))
        filesize = file.seek(0, os.SEEK_END)
        file.close()
        file = open(filename, 'rb')
        guid = ''.join(random.choices(string.ascii_letters +
                                      string.digits, k=40))
        chunkMetaData = {}
        chunkMetaData['FileName'] = filename
        chunkMetaData['TotalCount'] = '1'
        chunkMetaData['FileSize'] = filesize
        chunkMetaData['Index'] = 0
        chunkMetaData['FileGuid'] = guid
        chunkMetaData = json.dumps(chunkMetaData)
        # fileidUrl = leapUrl + '/api/datasets/generate/fileid?org=leo1311'
        # response = requests.get(url=fileidUrl, headers=fsHeaders, verify=False, proxies=proxyDict)
        # fileid = response.text
        uploadUrl = leapUrl + '/api/datasets/attachmentupload?org=LateralOffer'
        response = requests.post(url=uploadUrl, headers=fsHeaders, files=dict(file=file, chunkMetadata=chunkMetaData),
                                 verify=False, proxies=proxyDict)
        print('Attachment upload to fileserver response: ' + str(response.status_code))
        file.close()
        logger.info('Doc File Name :' + os.getcwd() + '\\' + filename)

        if attach['extension'] == '.zip':

            # filename=os.getcwd() + '\\' + filename

            logger.info('Zip file process started : ' + str(filename) + ' - ' + str(type(filename)))

            with ZipFile(filename, 'r') as f:

                fileNames = f.infolist()

                for elem in fileNames:

                    if '.docx' in elem.filename and '.doc' in elem.filename:
                        doc_Count += 1

                        f.extract(str(elem.filename), path=os.getcwd())

                        # destination = os.getcwd() + '\\' + elem.filename

                        destination = os.path.join(os.getcwd(), elem.filename)

                        my_dict['Doc File Path'] = destination
                        logger.info('Zip file process completed successfully ' + destination)

            f.close()

            logger.info('Zip file process completed successfully')

        elif attach['extension'] == '.docx' and '.doc' in filename:

            logger.info('Attached document process started')

            my_dict.update({'Doc File Path': os.path.join(os.getcwd(), str(filename))})

            logger.info('Attached document process completed successfully')

            doc_Count += 1
        if (response.status_code == 200):
            attachmentLink['filename'] = filename
            attachmentLink['fileid'] = response.json()['data']['fileId']
            # print(filename)
            # print(response.json()['data']['fileName'])
            # print(response.json()['data']['fileId'])
            attachmentLink['filesize'] = str(filesize)
            attachmentLink['bucket'] = projectName
            attachmentLink['uploadedOn'] = response.json()['data']['uploadedOn']
            if ('contentID' in attach):
                attachmentLink['contentID'] = attach['contentID']
        else:
            attachmentLink['filename'] = filename
            attachmentLink['fileid'] = 'Failed'
            attachmentLink['filesize'] = str(filesize)
            attachmentLink['bucket'] = projectName
            attachmentLink['contentID'] = attach['contentID']
            attachmentLink['uploadedOn'] = ''

        attachmentLinks.append(attachmentLink)
        print(attachmentLinks)

    timeStampafter_ParseAttachment = datetime.now()
    print('Time Stamp after ODS Data Extraction  ', str(timeStampafter_ParseAttachment))
    print('Time taken for ODS Data Extraction ', str(timeStampafter_ParseAttachment - timeStampBefore_ParseAttachment))
    return attachmentLinks


projectName = params['projectName']
logger.info(projectName)
leapDS = params['Authenticate']
leapDSdict = json.loads(leapDS)
connectionDetails = json.loads(leapDSdict['datasource']['connectionDetails'])
leapUrl = connectionDetails['Url']
attributes = json.loads(leapDSdict['attributes'])
authUrl = leapUrl + attributes['Url']
payload = attributes['Body']
# response = requests.post(url=authUrl, data=payload, verify=False, proxies=proxyDict)
# if response.status_code != 200:
#     logger.error('Failed to authenticate leap user. Status: ' + str(response.status_code))
#     exit()
# bearer = response.json()['id_token']
# print('User authenticated')
spineDS = params['Emails']

print(spineDS)
spineDSdict = json.loads(spineDS)
connectionDetails = spineDSdict['datasource']['connectionDetails']
connectionDetails = json.loads(connectionDetails)
spineUrl = connectionDetails['Url']
print('Spine URL: ' + spineUrl)
attributes = json.loads(spineDSdict['attributes'])
getMailsUrl = spineUrl + attributes['Url']
attributes = json.loads(spineDSdict['attributes'])
parameters = attributes['QueryParams']
queryParams = {}
for param in parameters:
    queryParams[param['key']] = param['value']
    # queryParams[param['key']] = '15'
'''
lastExecutionTime = params['Last_Successful_Execution_Time']
print(lastExecutionTime)
if lastExecutionTime != '-1':
    lastExecutionTime = (datetime.strptime(lastExecutionTime, '%a, %d %b %Y %H:%M:%S').astimezone(pytz.utc) - timedelta(
        minutes=10)).strftime('%Y-%m-%d %H:%M:%S')
    queryParams['ProcessedAfterTimestamp'] = lastExecutionTime
    logger.info(lastExecutionTime)
    # queryParams['MailID'] = '538'
    #queryParams['ProcessedAfterTimestamp'] = '2024-08-22+09%1A10%3A03'
'''
queryParams = urllib.parse.urlencode(queryParams)
if (queryParams):
    getMailsUrl += '?' + queryParams
print('Get mails URL: ' + getMailsUrl)
response = requests.get(getMailsUrl, verify=False, proxies=proxyDict)
if response.status_code != 200:
    logger.error('Failed to read emails. Status: ' + str(response.status_code))
    exit()
data = response.json()
print('Records Fetched {0}'.format(len(data)))
resultCount = len(data)
# convert json to spark dataframe
jsonStringArray = []
# payload = {'contents': jsonStringArray, 'comments': 'Pipeline assignment'}
headers = {'Content-type': 'application/json',
           'Accept': 'application/json, text/plain, */*', 'access-token': 'aec127c2-c984-33f6-9a3a-355xd1dof097'}
fsHeaders = {}
insertUrl = leapUrl + '/api/service/startProcess/' + projectName + '/{0}/email'
addEmailUrl = leapUrl + '/api/email'
getProjectUrl = leapUrl + '/api/service/projects/get/' + projectName
messageUrl = leapUrl + '/api/inbox/triggerMessage/{0}/mailReceived'
addEmailUrlNew = leapUrl + '/api/service/addEmail/' + projectName + '/{0}/mailReceived'
response = requests.get(url=getProjectUrl, headers=headers, verify=False, proxies=proxyDict)
if response.status_code != 200:
    logger.error('Project not found!')
    exit()
projectObj = response.json()
project = {}
project['id'] = response.json()['id']
# print('Printing the project response : ' + str(project ))
# print('Printing the data response: ' + str(data ))
# project.pop('createdDate')
spineToLeap = {'from': 'requested_by', 'subject': 'short_description', 'bodyText': 'task_details',
               'mailboxName': 'MailBoxId', 'importance': 'Priorities'}
# date.today().strftime('%d/%m/%Y'):'CreatedDate'
# print('Printing the spineToLeap response: ' + str(spineToLeap ))
spineToUsm = {'from': 'email_from', 'subject': 'email_subject', 'body': 'email_body', 'to': 'email_to',
              'cc': 'email_cc', 'bcc': 'email_bcc', 'noOfAttachments': 'attachment_count',
              'mailType': 'mail_type', 'messageID': 'mailId', 'mailboxName': 'mail_box_name'}

keywordToProcess = {'hr': 'lateraloffer', 'InfyHR': 'lateraloffer'}
for row in data:
    # logger.info(' Mail Id ',str(row['messageID']))
    print(row['messageID'])
    timeStampBeforeDataExtraction = datetime.now()
    print('start time for this mail id ', str(timeStampBeforeDataExtraction))

    subject = row['subject']
    print(subject)

    if row['mailboxName'].lower() == 'infyhrread':
        print('Parsing email: ' + subject)
        print(row['messageID'])

        usmEmail = {}

        for key in spineToUsm.keys():
            usmEmail[spineToUsm[key]] = row[key]

        usmEmail['projectId'] = project

        usmEmail['last_updated_date'] = datetime.now(timezone.utc).isoformat()

        usmEmail['is_processed'] = True

        usmEmail['sent_date'] = datetime.now(timezone.utc).isoformat()

        usmEmail['is_active'] = True

        usmEmail['isRead'] = False

        usmEmail['is_incoming'] = True

        # print(usmEmail)
        # usmEmail['caseId'] = case_id
        logger.info('ccount aattachments : ' + str(len(row['attachments'])))
        if (len(row['attachments']) > 0):
            usmEmail['attachments'] = json.dumps(parse_attachments(row['attachments'], row['messageID']))
        # else:
        #     usmEmail['attachments'] = '[]'

        case_id = ''

        caseIdRegex = re.compile(r'\w{2,4}:\d{8}')

        if caseIdRegex.search(subject) != None:

            case_id = caseIdRegex.search(subject).group()

            print('Email received for: ' + case_id)

            usmEmail['caseId'] = case_id

            data = {}

            '''
            # data['message'] = row['body'].split('From:')[0]
            if(row['closingCounter']>1):
               data['status'] = 'completed' 

            p = re.compile(r'clarification')
            if (p.search(subject.lower())):
                data['status_'] = 'clarificationreceived'
            '''
            data['status_'] = 'clarificationreceived'

            data['priority'] = 'low'

            data['requireClarification'] = False

            data['email'] = usmEmail

            messageUrlFmt = addEmailUrlNew.format(case_id)
            print(messageUrlFmt)
            response = requests.post(url=messageUrlFmt, data=json.dumps(data), headers=headers, verify=False,
                                     proxies=proxyDict)

            if (response.status_code == 200):
                print('Email already parsed')

                continue

            if (response.status_code != 201):
                logger.error('Error adding email. Status: ' + str(response.status_code))

                pipelineStatus = 'Error'

                continue

            print('Email successfully added')



        else:

            for keyword in keywordToProcess.keys():

                if row['importance'] == 'Unknown':
                    row['importance'] = 'normal'

                if keyword in subject.lower():

                    processName = keywordToProcess[keyword]

                    insertUrlFmt = insertUrl.format(processName)

                    print(insertUrlFmt)

                    # for attkeyword in keywordToProcess.keys():

                    #  if attkeyword in subject.lower():

                    #  row['TicketType'] = keywordToProcess[attkeyword]

                    for key in spineToLeap.keys():
                        row[spineToLeap[key]] = row[key]

                    row['manualCreator'] = ''

                    row['processName'] = processName

                    row['CreatedDate'] = date.today().strftime('%Y/%m/%d')

                    row['TicketType'] = 'infy'

                    row['uploadControl'] = '[]'

                    row['email'] = usmEmail
                    # row['ticketCategory']=category
                    if doc_Count == 1:
                        try:
                            ODS_Extraction(my_dict)
                            print(final_dict)
                            row['candidateName'] = final_dict['Candidate Name']
                            row['accountName'] = final_dict['Account Name']
                            row['rwi'] = final_dict['Restart With Infosys']
                            row['src'] = final_dict['Source']
                            row['tech'] = final_dict['Technology']
                            row['candidateId'] = final_dict['Candidate ID']
                            row['indent'] = final_dict['Indent']
                            row['roleDesignation'] = final_dict['Role Designation']
                            row['jobLevel'] = final_dict['Job Level']
                            row['personalLevel'] = final_dict['personal level']
                            row['jobSubLevel'] = final_dict['Job sub-Level']
                            row['personalSubLevel'] = final_dict['personal sub-Level']
                            row['category'] = final_dict['Category']
                            row['practiceUnit'] = final_dict['Unit']
                            row['unit'] = final_dict['Practice Unit']
                            row['locationOfPosting'] = final_dict['Location of Posting']
                            row['establishmentType'] = final_dict['Establishment']
                            row['hiringType'] = final_dict['Hiring Type']
                            row['remoteWorkFromHome'] = final_dict['Remote work from home']
                            row['permanentWorkFromHome'] = final_dict['Permanent work from home']
                            row['remoteWorkLocation'] = final_dict['Remote work Location']
                            row['partTimeEmployee'] = final_dict['partTimeEmployee']
                            row['educationCategory'] = final_dict['Education Category']
                            row['experience'] = final_dict['Experience (Month)']
                            try:
                                final_dict['Date of joining'] = datetime.strptime(str(final_dict['Date of joining']),
                                                                                  '%m-%d-%Y')  # converting to date format
                                row['dateOfJoining'] = final_dict['Date of joining'].strftime('%Y/%m/%d %H:%M:%S')
                            except Exception as e:
                                logger.error('Error in dateformat Date of joining ' + str(e))
                            try:
                                final_dict['End Date'] = datetime.strptime(str(final_dict['End Date']),
                                                                           '%m-%d-%Y')  # converting to date format
                                row['endDate'] = final_dict['End Date'].strftime('%Y/%m/%d %H:%M:%S')
                            except Exception as e:
                                logger.error('Error in dateformat End date ' + str(e))

                            row['salaryProposed'] = final_dict['Salary Proposed (in Lacs)']
                            row['joiningBonus'] = final_dict['Joining Bonus']
                            row['t50'] = final_dict['T50']
                            row['ae'] = final_dict['AE']
                            row['stretch'] = final_dict['Stretch %']
                            row['ccEmployeeName'] = final_dict['CC Employee Name']
                            row['recuriterIdApproval'] = final_dict['Recruiter ID - Contact Details']
                            row['recruiterManagerId'] = final_dict['Recruiter Manager ID']
                            row['recuriterNameApproval'] = final_dict['Recuriter Name - Contact Details']
                            row['recuriterManagerName'] = final_dict['Recuriter Manager Name']
                            row['employeeNameOne'] = final_dict['Employee Name1']
                            row['employeeIdOne'] = final_dict['Employee ID1']
                            row['employeeRoleOne'] = final_dict['Employee Role1']
                            row['employeeNameTwo'] = final_dict['Employee Name2']
                            row['employeeIdTwo'] = final_dict['Employee ID2']
                            row['employeeRoleTwo'] = final_dict['Employee Role2']
                            row['employeeNameThree'] = final_dict['Employee Name3']
                            row['employeeIdThree'] = final_dict['Employee ID3']
                            row['employeeRoleThree'] = final_dict['Employee Role3']
                            row['digitalOffering'] = final_dict['NSO']
                            dtDigitalOffer = final_dict['NSO']
                            dtPuCode = final_dict['Hiring Type']
                            dtUnit = final_dict['Practice Unit']

                            row1 = {'dtDigitalOffer': {dtDigitalOffer}, 'dtPuCode': {dtPuCode}, 'dtUnit': {dtUnit}}
                            digitalTag = None
                            # df_filtered = df.query('dtDigitalOffer == {dtDigitalOffer} and dtPuCode=={dtPuCode} and dtUnit=={dtUnit}'.format(dtDigitalOffer=dtDigitalOffer,dtPuCode=row['dtPuCode'],dtUnit=row['dtUnit']))
                            df_filtered = df.query(
                                'dtDigitalOffer == @dtDigitalOffer and dtPuCode==@dtPuCode and dtUnit==@dtUnit'.format(
                                    **row1))
                            # df_filtered = df.query('dtDigitalOffer == {dtDigitalOffer} '.format(dtDigitalOffer)
                            if len(df_filtered) == 1:
                                find_index = df_filtered['dtDigitalTag']
                                digitalTag = df_filtered.loc[find_index.index.tolist()[0], 'dtDigitalTag']
                                # print(find_index.index.tolist()[0])
                            else:
                                digitalTag = 'NA'
                            print('Digital tag : ' + digitalTag)

                            row['digitalTag'] = digitalTag
                        except Exception as e:
                            final_dict = {'Candidate ID': '', 'Indent': '', 'Role Designation': '', 'Job Level': '',
                                          'personal level': '', 'Job sub-Level': '', 'personal sub-Level': '',
                                          'Category': '', 'Practice Unit': '', 'Location of Posting': '',
                                          'Establishment': '', 'Hiring Type': '', 'Work From Home': '',
                                          'Remote work from home': '', 'Permanent work from home': '',
                                          'Remote work Location': '', 'Education Category': '',
                                          'Experience (Month)': '', 'Salary Proposed (in Lacs)': None,
                                          'Joining Bonus': None, 'T50': '', 'AE': None, 'Stretch %': None,
                                          'CC Employee Name': '', 'Recruiter ID - Contact Details': '',
                                          'Recruiter Manager ID': '', 'End Date': '', 'Date of joining': '',
                                          'Recuriter Manager Name': '', 'Recuriter Name - Contact Details': ''}
                            logger.error('Error occurred in values mapping ' + str(e))
                    elif not any(final_dict.values()) or doc_Count == 0 or doc_Count > 1:
                        final_dict = {'Candidate ID': '', 'Indent': '', 'Role Designation': '', 'Job Level': '',
                                      'personal level': '', 'Job sub-Level': '', 'personal sub-Level': '',
                                      'Category': '', 'Practice Unit': '', 'Location of Posting': '',
                                      'Establishment': '', 'Hiring Type': '', 'Work From Home': '',
                                      'Remote work from home': '', 'Permanent work from home': '',
                                      'Remote work Location': '', 'Education Category': '',
                                      'Experience (Month)': '', 'Salary Proposed (in Lacs)': None,
                                      'Joining Bonus': None, 'T50': '', 'AE': None, 'Stretch %': None,
                                      'CC Employee Name': '', 'Recruiter ID - Contact Details': '',
                                      'Recruiter Manager ID': '', 'End Date': '', 'Date of joining': '',
                                      'Recuriter Manager Name': '', 'Recuriter Name - Contact Details': ''}
                    logger.info(
                        'Document is not attached or attached more than 1 documents. Hence the ODS Extraction fields will remain empty or none')

                row.pop('body')

                row.pop('bodyText')

                row.pop('task_details')

                print(json.dumps(row))
                print(final_dict['Practice Unit'])
                print(lst_UnitsForCaseCreation)

                timeStampBeforePostCall = datetime.now()
                print('Time Stamp before post call  ', str(timeStampBeforePostCall))

                response = requests.post(url=insertUrlFmt, data=json.dumps(row), headers=headers, verify=False,
                                         proxies=proxyDict)

                timeStampAfterPostCall = datetime.now()
                print('Time Stamp after post call   ', str(timeStampAfterPostCall))
                print('Time taken for this mail id, post call ', str(timeStampAfterPostCall - timeStampBeforePostCall))

                if (response.status_code == 200):
                    print('Email already parsed')

                    break

                if (response.status_code != 201):
                    logger.error('Error creating case. Status: ' + str(response.status_code))

                    logger.error(response.text)

                    pipelineStatus = 'Error'

                    break

                case_id = response.text.split(':', 1)[1]

                print('Case created: ' + case_id)

                break

                '''
                if final_dict['Practice Unit'] in lst_UnitsForCaseCreation and 're:' not in subject.lower():
                    print('Practice Unit value is there in config table')

                else:
                    print('Practice Unit value is not there in config table')
                '''

timeStampAfterDataExtraction = datetime.now()
print('Time Stamp after case creation   ', str(timeStampAfterDataExtraction))
print('Time taken for this mail id, case creation ', str(timeStampAfterDataExtraction - timeStampBeforeDataExtraction))

print(pipelineStatus)
# print('Completed')








































































































































